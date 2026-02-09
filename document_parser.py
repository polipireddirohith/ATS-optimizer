"""
Document Parser for ATS Engine

Handles parsing of resume files in multiple formats:
- PDF
- DOCX (Microsoft Word)
- TXT (Plain text)
"""

import os
from typing import Optional


class DocumentParser:
    """Parse resume documents from various formats"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_file(self, file_path: str) -> str:
        """
        Parse a resume file and extract text
        
        Args:
            file_path: Path to the resume file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            FileNotFoundError: If file doesn't exist
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_ext}. Supported formats: {', '.join(self.supported_formats)}")
        
        if file_ext == '.pdf':
            return self._parse_pdf(file_path)
        elif file_ext == '.docx':
            return self._parse_docx(file_path)
        elif file_ext == '.txt':
            return self._parse_txt(file_path)
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file"""
        try:
            import PyPDF2
            
            text = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text.append(page.extract_text())
            
            return '\n'.join(text)
        except ImportError:
            print("PyPDF2 not installed. Attempting alternative method...")
            try:
                import pdfplumber
                
                text = []
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text.append(page.extract_text())
                
                return '\n'.join(text)
            except ImportError:
                raise ImportError(
                    "PDF parsing requires PyPDF2 or pdfplumber. "
                    "Install with: pip install PyPDF2 or pip install pdfplumber"
                )
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text.append(cell.text)
            
            return '\n'.join(text)
        except ImportError:
            raise ImportError(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            )
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                return file.read()
    
    def parse_text(self, text: str) -> str:
        """
        Parse text directly (for when text is already extracted)
        
        Args:
            text: Resume text
            
        Returns:
            Cleaned text
        """
        return text.strip()


def main():
    """Test document parser"""
    parser = DocumentParser()
    print("Document Parser initialized")
    print(f"Supported formats: {', '.join(parser.supported_formats)}")


if __name__ == "__main__":
    main()
